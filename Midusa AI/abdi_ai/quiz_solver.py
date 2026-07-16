import os
import base64
import logging
import tempfile
from groq import Groq
from docx import Document

logger = logging.getLogger(__name__)

client = Groq(api_key=os.getenv("GROQ_API_KEY"))

def extract_text_from_pdf(file_path: str) -> str:
    import fitz
    doc = fitz.open(file_path)
    text = ""
    for page in doc:
        text += page.get_text()
    doc.close()
    return text

def encode_image(file_path: str) -> str:
    with open(file_path, "rb") as f:
        return base64.b64encode(f.read()).decode("utf-8")

def solve_quiz_text(content: str) -> str:
    prompt = f"""Solve this quiz/exam. Format your response like this:

Question 1: [question text]
Answer 1: [correct answer]
Explanation 1: [brief explanation]

Question 2: [question text]
Answer 2: [correct answer]
Explanation 2: [brief explanation]

Quiz content:
{content}"""
    
    response = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.1,
    )
    return response.choices[0].message.content

def solve_quiz_image(file_path: str) -> str:
    image_data = encode_image(file_path)
    response = client.chat.completions.create(
        model="meta-llama/llama-4-scout-17b-16e-instruct",
        messages=[{
            "role": "user",
            "content": [
                {"type": "text", "text": "Solve this quiz/exam shown in the image. Format your response:\n\nQuestion 1: [question]\nAnswer 1: [answer]\nExplanation 1: [explanation]\n\nQuestion 2: [question]\nAnswer 2: [answer]\nExplanation 2: [explanation]"},
                {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{image_data}"}}
            ]
        }],
        temperature=0.1,
    )
    return response.choices[0].message.content

def create_solved_document(original_name: str, solution_text: str) -> str:
    doc = Document()
    title = doc.add_heading(f"Solved Quiz", 0)
    doc.add_paragraph(f"Original file: {original_name}")
    doc.add_paragraph("─" * 50)

    for line in solution_text.split("\n"):
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("**") and stripped.endswith("**"):
            doc.add_heading(stripped.strip("*"), 2)
        elif any(stripped.lower().startswith(w) for w in ["question", "q:", "q.", "answer", "a:", "a.", "explanation", "note"]):
            p = doc.add_paragraph()
            runner = p.add_run(stripped)
            runner.bold = True
        else:
            doc.add_paragraph(stripped)

    base = os.path.splitext(original_name)[0]
    output_path = os.path.join(tempfile.gettempdir(), f"solved_{base}.docx")
    doc.save(output_path)
    return output_path

async def solve_quiz_file(file_path: str, file_name: str) -> str:
    logger.info(f"Solving quiz: {file_name}")
    ext = file_name.lower()
    
    if ext.endswith(".pdf"):
        content = extract_text_from_pdf(file_path)
        if not content.strip():
            raise ValueError("Could not extract text from PDF")
        solution = solve_quiz_text(content[:10000])
    elif ext.endswith((".png", ".jpg", ".jpeg", ".gif", ".webp")):
        solution = solve_quiz_image(file_path)
    else:
        raise ValueError("Unsupported format. Send PDF or image.")
    
    output_path = create_solved_document(file_name, solution)
    return output_path
